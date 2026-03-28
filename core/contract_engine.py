import docx
import re

class PentagonScoutAuditorV50000:
    def __init__(self, doc_path):
        self.doc = docx.Document(doc_path)
        self.segments = []
        for p in self.doc.paragraphs:
            if p.text.strip(): self.segments.append(p.text.strip())
        
        self.table_cells = []
        for table in self.doc.tables:
            for row in table.rows:
                row_data = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                row_str = " | ".join(row_data)
                self.segments.append(row_str)
                self.table_cells.append((row_data, row_str))
                
        self.full_content = "\n".join(self.segments)
        self.total_chars = len(self.full_content)

    def extract_dossier_v50(self):
        res = {"total": "Chưa rõ", "vat": "Chưa rõ", "exp_count": "12", "exp_cost": "Chưa rõ", "other": "Chưa rõ", "duration": "Chưa rõ"}
        for cells, row_str in self.table_cells:
            low = row_str.lower()
            nums = re.findall(r"\d{1,3}(?:[.,]\d{3})*", row_str)
            if not nums: continue
            if "chuyên gia" in low or "thù lao" in low: res["exp_cost"] = nums[-1]
            if "chi phí khác" in low: res["other"] = nums[-1]
            if "thuế" in low or "vat" in low:
                v_p = re.search(r"(\d+)\s?%", low)
                res["vat"] = v_p.group(1) + "%" if v_p else nums[-1]
        m_total = re.findall(r"(\d{1,3}(?:[.,]\d{3})+(?:\s?(?:vnđ|vnd|đồng|đ|usd|\$))?)", self.full_content, re.I)
        if m_total: res["total"] = m_total[-1]
        d_match = re.search(r"(\d+)\s?tháng", self.full_content, re.I)
        if d_match: res["duration"] = d_match.group(1) + " tháng"
        return res

    def deep_pentagon_audit_v50(self):
        audit_plan = [
            {"Nhóm": "⚖️ PHÁP LÝ", "Mục": "Phạt vi phạm", "Key": r"phạt|vi phạm|penalty", "Advice": "Kiểm tra mức phạt tối đa 8%."},
            {"Nhóm": "⚖️ PHÁP LÝ", "Mục": "Tranh chấp", "Key": r"trọng tài|tranh chấp", "Advice": "Ưu tiên Trọng tài thương mại."},
            {"Nhóm": "💸 TÀI CHÍNH", "Mục": "Tạm ứng & Bảo lãnh", "Key": r"tạm ứng|bảo lãnh", "Advice": "Đảm bảo thời hạn bảo lãnh bao trùm."},
            {"Nhóm": "💸 TÀI CHÍNH", "Mục": "Thanh toán", "Key": r"kỳ thanh toán|đợt thanh toán", "Advice": "Kiểm soát dòng tiền chuyên gia."},
            {"Nhóm": "🛠️ KỸ THUẬT", "Mục": "Phạm vi & Tiêu chuẩn", "Key": r"phạm vi công việc|tiêu chuẩn", "Advice": "Tránh nới rộng phạm vi ngoài dự báo."},
            {"Nhóm": "👥 NHÂN SỰ", "Mục": "Thay đổi chuyên gia", "Key": r"thay đổi nhân sự|thay thế chuyên gia", "Advice": "Thủ tục thay đổi phải đơn giản."},
            {"Nhóm": "👥 NHÂN SỰ", "Mục": "Bằng cấp", "Key": r"bằng cấp|chứng chỉ hành nghề", "Advice": "Đối soát chứng chỉ 12 chuyên gia."},
            {"Nhóm": "📅 TIẾN ĐỘ", "Mục": "Milestone", "Key": r"mốc tiến độ|giai đoạn|milestone", "Advice": "Làm rõ mốc hoàn thành báo cáo đợt 1, 2..."},
            {"Nhóm": "📅 TIẾN ĐỘ", "Mục": "EOT", "Key": r"gia hạn|kéo dài thời gian|extension of time", "Advice": "Lưu ý quyền được gia hạn do bất khả kháng."}
        ]
        detailed_report = []
        summary_stats = {"⚖️ PHÁP LÝ": 0, "💸 TÀI CHÍNH": 0, "🛠️ KỸ THUẬT": 0, "👥 NHÂN SỰ": 0, "📅 TIẾN ĐỘ": 0}
        for p in audit_plan:
            match = re.search(p["Key"], self.full_content, re.I)
            text, desc, status = "⚠️ Bỏ trống", "Cần bổ sung điều khoản.", "⚠️ Thiếu sót"
            if match:
                summary_stats[p["Nhóm"]] += 1
                status = "✅ Đã có"
                start = max(0, match.start() - 150)
                end = min(self.total_chars, match.start() + 450)
                text = "..." + self.full_content[start:end] + "..."
                desc = f"Hợp đồng có quy định về {p['Mục']}."
            detailed_report.append({
                "Nhóm": p["Nhóm"], "Hạng mục": p["Mục"], "Trạng thái": status,
                "Trích dẫn nguyên văn (Evidence)": text,
                "AI Giải mã nội dung (Description)": desc,
                "Khuyến nghị của AI (Advice)": p["Advice"]
            })
        return summary_stats, detailed_report

    def run_pentagon_audit(self):
        s, d = self.deep_pentagon_audit_v50()
        return {
            "dossier": self.extract_dossier_v50(),
            "summary_stats": s,
            "detailed_audit": d,
            "total_chars": self.total_chars
        }

def analyze_contract_v50000(input_path):
    return PentagonScoutAuditorV50000(input_path).run_pentagon_audit()
